from docx import Document

doc = Document()
doc.add_heading("Minutes of Meeting", level=1)

with open("outputs/MoM.txt", "r", encoding="utf-8") as f:
    doc.add_paragraph(f.read())

doc.save("outputs/MoM.docx")
print("Word file created")
